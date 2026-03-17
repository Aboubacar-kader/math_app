"""
Service de traitement des documents uploadés.
Extrait le texte de différents formats de fichiers.
"""

import pdfplumber
from docx import Document
import io
import re
from typing import List, Dict, Any, Optional
import streamlit as st
from pathlib import Path
from config.settings import settings
from core.llm_manager import call_1minai

def _parse_vision_response(data: dict) -> str:
    """Extrait le texte d'une réponse JSON OpenAI vision (choices[0].message.content)."""
    try:
        return data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError, TypeError):
        return ""


def _call_vision(b64_image: str, prompt: str, mime: str = "image/png") -> str:
    """Appelle l'API OpenAI vision avec une image base64. Retourne le texte extrait."""
    import requests as _requests
    url = f"{settings.MIN_AI_BASE_URL.rstrip('/')}/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {settings.MIN_AI_API_KEY.strip()}",
    }
    payload = {
        "model": settings.MIN_AI_MODEL,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64_image}"}},
                ],
            }
        ],
        "max_tokens": 4096,
    }
    response = _requests.post(url, headers=headers, json=payload, timeout=90)
    response.raise_for_status()
    return _parse_vision_response(response.json())


class DocumentProcessor:
    """Processeur de documents multi-formats"""
    
    SUPPORTED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'image/png': 'image',
        'image/jpeg': 'image',
        'image/jpg': 'image',
        'image/x-png': 'image',
        'image/webp': 'image',
        'image/bmp': 'image',
        'image/gif': 'image',
        'image/tiff': 'image',
        'image/x-ms-bmp': 'image',
    }
    
    def __init__(self):
        self.upload_dir = settings.UPLOAD_DIR
    
    def process_files(self, uploaded_files: List[Any]) -> List[Dict[str, Any]]:
        """
        Traite une liste de fichiers uploadés.
        
        Args:
            uploaded_files: Liste des fichiers Streamlit
            
        Returns:
            Liste de dictionnaires contenant le texte et les métadonnées
        """
        processed_docs = []
        
        for file in uploaded_files:
            try:
                # Vérifier la taille
                if file.size > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
                    st.warning(f"⚠️ {file.name} dépasse {settings.MAX_FILE_SIZE_MB}MB")
                    continue
                
                # Extraire le texte selon le type
                text = self._extract_text(file)

                # Limite de sécurité sur le texte extrait (anti décompression bomb)
                MAX_EXTRACTED_CHARS = 500_000
                if text and len(text) > MAX_EXTRACTED_CHARS:
                    text = text[:MAX_EXTRACTED_CHARS]

                # Nom de fichier sûr (basename uniquement, caractères autorisés)
                safe_name = Path(file.name).name
                safe_name = re.sub(r'[^\w\.\-]', '_', safe_name)

                if text and text.strip():
                    processed_docs.append({
                        'text': text,
                        'metadata': {
                            'filename': safe_name,
                            'type': file.type,
                            'size': file.size
                        }
                    })

                    # Sauvegarder le fichier
                    self._save_file(file, safe_name)
                else:
                    st.warning(f"⚠️ Aucun texte extrait de {file.name}")

            except Exception as e:
                st.error(f"❌ Erreur avec {file.name}: traitement impossible")
        
        return processed_docs
    
    def _extract_text(self, file) -> str:
        """
        Extrait le texte d'un fichier selon son type.
        Fallback par extension si le MIME type n'est pas reconnu.
        """
        file_type = self.SUPPORTED_TYPES.get(file.type)

        # Fallback par extension si MIME type inconnu (ex : image/x-png sur certains navigateurs)
        if not file_type:
            ext = Path(file.name).suffix.lower()
            if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif'):
                file_type = 'image'
            elif ext == '.pdf':
                file_type = 'pdf'
            elif ext in ('.docx', '.doc'):
                file_type = 'docx'
            elif ext == '.txt':
                file_type = 'txt'

        if file_type == 'pdf':
            return self._extract_from_pdf(file)
        elif file_type == 'docx':
            return self._extract_from_docx(file)
        elif file_type == 'txt':
            return self._extract_from_txt(file)
        elif file_type == 'image':
            return self._extract_from_image(file)
        else:
            st.warning(f"⚠️ Type de fichier non supporté : {file.type} ({file.name})")
            return ""
    
    def _extract_from_pdf(self, file) -> str:
        """Extrait le texte d'un PDF via pdfplumber.
        Si une page ne contient pas de texte sélectionnable (PDF scanné),
        elle est envoyée à la vision IA pour OCR."""
        try:
            raw_bytes = file.read()
            file.seek(0)
            file_bytes = io.BytesIO(raw_bytes)
            text = ""

            with pdfplumber.open(file_bytes) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""

                    # Tableaux
                    tables = page.extract_tables()
                    table_text = ""
                    for table in tables:
                        for row in table:
                            row_cells = [str(cell or "").strip() for cell in row]
                            table_text += " | ".join(row_cells) + "\n"
                        table_text += "\n"

                    page_content = page_text
                    if table_text.strip():
                        page_content += "\n" + table_text

                    # Fallback vision si page vide (PDF scanné / image)
                    if not page_content.strip():
                        page_content = self._ocr_pdf_page(page, page_num + 1)

                    if page_content.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_content}"

            return text
        except Exception as e:
            st.error(f"Erreur PDF: {str(e)}")
            return ""

    def _ocr_pdf_page(self, page, page_num: int) -> str:
        """Convertit une page PDF en image et l'envoie à l'API OpenAI vision pour OCR."""
        try:
            import base64

            # Rendre la page en image (pdfplumber utilise pdfminer/Pillow)
            pil_image = page.to_image(resolution=150).original
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

            prompt = (
                "Transcris intégralement tout le texte et toutes les formules "
                "mathématiques de cette page de document. "
                "Conserve la structure : titres, numéros, tableaux, formules "
                "(notation standard : f(x) = ...). Ne résous rien, transcris uniquement."
            )
            return _call_vision(b64, prompt)
        except Exception as e:
            st.warning(f"⚠️ OCR page {page_num} échoué : {e}")
            return ""
    
    def _extract_from_docx(self, file) -> str:
        """Extrait le texte d'un document Word (paragraphes + tableaux)"""
        try:
            file_bytes = io.BytesIO(file.read())
            doc = Document(file_bytes)

            parts = []

            # Paragraphes
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(c for c in row_cells if c)
                    if row_text:
                        parts.append(row_text)

            return "\n\n".join(parts)
        except Exception as e:
            st.error(f"Erreur DOCX: {str(e)}")
            return ""
    
    def _extract_from_txt(self, file) -> str:
        """Extrait le texte d'un fichier texte"""
        try:
            return file.read().decode('utf-8')
        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            try:
                file.seek(0)
                return file.read().decode('latin-1')
            except:
                return ""
    
    def _extract_from_image(self, file) -> str:
        """Extraction via API OpenAI vision — aucun OCR local requis."""
        import base64

        file_bytes = file.read()
        file.seek(0)

        # Déterminer le MIME type pour l'URL data
        mime = file.type if file.type and file.type.startswith('image/') else 'image/png'
        b64_image = base64.b64encode(file_bytes).decode('utf-8')

        try:
            prompt = (
                "Transcris intégralement tout le texte et toutes les formules "
                "mathématiques de cette image. "
                "Garde la structure exacte de l'exercice : titres, numéros de questions, "
                "tableaux, formules (notation standard : f(x) = ...). "
                "Ne résous rien, transcris uniquement."
            )
            text = _call_vision(b64_image, prompt, mime=mime)
            if text:
                return text
        except Exception as e:
            st.warning(f"⚠️ Extraction vision échouée : {e}")

        return "[Image - extraction non disponible]"
    
    def _save_file(self, file, safe_name: str = None):
        """Sauvegarde le fichier uploadé avec protection contre le path traversal."""
        try:
            name = safe_name or Path(file.name).name
            upload_dir_resolved = self.upload_dir.resolve()
            file_path = (upload_dir_resolved / name).resolve()

            # Vérification anti path traversal
            if not str(file_path).startswith(str(upload_dir_resolved)):
                st.warning(f"Nom de fichier invalide : {file.name}")
                return

            with open(file_path, 'wb') as f:
                f.write(file.getbuffer())
        except Exception:
            st.warning("Sauvegarde du fichier échouée.")
    
# Instance globale
document_processor = DocumentProcessor()